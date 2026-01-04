import docx
import os

def create_sample():
    doc = docx.Document()
    doc.add_heading("Neon Nights", 0)
    doc.add_paragraph("#synthwave")
    doc.add_paragraph("#cyberpunk")
    
    doc.add_paragraph("[Verse 1]")
    doc.add_paragraph("Walking down the street")
    doc.add_paragraph("Lights are flashing")
    
    doc.add_paragraph("[Chorus]")
    doc.add_paragraph("Neon nights, hold me tight")
    
    os.makedirs("vault/raw", exist_ok=True)
    doc.save("vault/raw/neon-nights.docx")
    print("Created vault/raw/neon-nights.docx")

if __name__ == "__main__":
    create_sample()